"""Generación de informes Word (.docx) del POA a partir del JSON del dashboard.

Formatos disponibles:
  institucional  -> Informe institucional general (todas las categorías + proyectos)
  administrativas-> Informe anual de las unidades administrativas
  academicas     -> Informe anual de las unidades académicas
  carreras       -> Informe anual de las carreras técnicas y tecnológicas
"""

from __future__ import annotations

import datetime
import re
from io import BytesIO

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

INSTITUCION = (
    "INSTITUTO SUPERIOR TECNOLÓGICO DEL AZUAY CON CONDICIÓN DE SUPERIOR UNIVERSITARIO"
)

AZUL = RGBColor(0x1F, 0x3C, 0x88)
GRIS = "D9E2F3"

REPORTES = {
    "institucional": {
        "categorias": ["Unidad Administrativa", "Unidad Académica", "Carrera"],
        "titulo": "SEGUIMIENTO AL PLAN OPERATIVO ANUAL INSTITUCIONAL",
        "subtitulo": "INFORME INSTITUCIONAL GENERAL",
        "sujeto": "las unidades administrativas, académicas y carreras",
        "archivo": "informe_poa_institucional",
        "proyectos": True,
    },
    "administrativas": {
        "categorias": ["Unidad Administrativa"],
        "titulo": "SEGUIMIENTO AL PLAN OPERATIVO DE LAS UNIDADES ADMINISTRATIVAS",
        "subtitulo": "INFORME ANUAL",
        "sujeto": "las unidades administrativas",
        "archivo": "informe_poa_unidades_administrativas",
        "proyectos": False,
    },
    "academicas": {
        "categorias": ["Unidad Académica"],
        "titulo": "SEGUIMIENTO AL PLAN OPERATIVO DE LAS UNIDADES ACADÉMICAS",
        "subtitulo": "INFORME ANUAL",
        "sujeto": "las unidades académicas",
        "archivo": "informe_poa_unidades_academicas",
        "proyectos": False,
    },
    "carreras": {
        "categorias": ["Carrera"],
        "titulo": (
            "SEGUIMIENTO AL PLAN OPERATIVO DE LAS CARRERAS TÉCNICAS Y TECNOLÓGICAS"
        ),
        "subtitulo": "INFORME ANUAL",
        "sujeto": "las carreras técnicas y tecnológicas",
        "archivo": "informe_poa_carreras",
        "proyectos": False,
    },
}


# --------------------------------------------------------------------------- utilidades


def pct(valor) -> str:
    if valor is None:
        return "N/D"
    return f"{round(float(valor) * 100, 2):g}%"


def _sombrear(celda, color=GRIS):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color)
    celda._tc.get_or_add_tcPr().append(shd)


def _celda(celda, texto, negrita=False, centrado=False, tam=9):
    celda.text = ""
    p = celda.paragraphs[0]
    if centrado:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(texto))
    run.bold = negrita
    run.font.size = Pt(tam)


def _tabla(doc, encabezados, filas, anchos=None):
    tabla = doc.add_table(rows=1, cols=len(encabezados))
    tabla.style = "Table Grid"
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, texto in enumerate(encabezados):
        _celda(tabla.rows[0].cells[i], texto, negrita=True, centrado=True)
        _sombrear(tabla.rows[0].cells[i])
    for fila in filas:
        celdas = tabla.add_row().cells
        for i, valor in enumerate(fila):
            _celda(celdas[i], valor, centrado=(i == 0 or i == len(fila) - 1))
    if anchos:
        for fila in tabla.rows:
            for i, ancho in enumerate(anchos):
                fila.cells[i].width = Inches(ancho)
    return tabla


def _titulo(doc, texto, nivel=1):
    h = doc.add_heading(texto, level=nivel)
    for run in h.runs:
        run.font.color.rgb = AZUL
    return h


def _parrafo(doc, texto, negrita=False, tam=11, centrado=False):
    p = doc.add_paragraph()
    if centrado:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texto)
    run.bold = negrita
    run.font.size = Pt(tam)
    return p


def _grafico_barras(titulo, etiquetas, valores):
    """Devuelve un PNG en memoria o None si matplotlib no está disponible."""
    try:
        import matplotlib

        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
    except Exception:  # noqa: BLE001
        return None

    alto = max(2.4, 0.32 * len(etiquetas) + 1)
    fig, ax = plt.subplots(figsize=(8.5, alto))
    pos = range(len(etiquetas))
    colores = ["#2f9e5f" if v >= 0.8 else "#e0a12b" if v >= 0.4 else "#c0392b" for v in valores]
    ax.barh(list(pos), [v * 100 for v in valores], color=colores)
    ax.set_yticks(list(pos))
    ax.set_yticklabels([e[:45] for e in etiquetas], fontsize=8)
    ax.invert_yaxis()
    ax.set_xlim(0, 100)
    ax.set_xlabel("% de avance")
    ax.set_title(titulo, fontsize=10, fontweight="bold")
    for i, v in enumerate(valores):
        ax.text(min(v * 100 + 1, 95), i, f"{round(v * 100, 1):g}%", va="center", fontsize=7)
    ax.grid(axis="x", linestyle=":", alpha=0.5)
    fig.tight_layout()
    buffer = BytesIO()
    fig.savefig(buffer, format="png", dpi=150)
    plt.close(fig)
    buffer.seek(0)
    return buffer


def _estados(actividades):
    con = [a for a in actividades if a.get("avance") is not None]
    return {
        "total": len(actividades),
        "completadas": len([a for a in con if a["avance"] >= 1]),
        "proceso": len([a for a in con if 0 < a["avance"] < 1]),
        "sin_iniciar": len([a for a in con if a["avance"] == 0]),
    }


def _actividades_validas(unidad):
    """Descarta filas que no son actividades reales (separadores, números sueltos)."""
    limpias = []
    for a in unidad.get("actividades") or []:
        tarea = (a.get("tarea") or "").strip()
        if not tarea or len(tarea) < 5:
            continue
        if re.match(r"^objetivo\s+(operativo|estrat)", tarea, re.I):
            continue
        if re.match(r"^[\d\s.,%/-]+$", tarea):
            continue
        limpias.append(a)
    return limpias


def _nombre_unidad(u):
    return (u.get("titulo") or u.get("nombre") or "").strip() or u.get("nombre", "")


# --------------------------------------------------------------------------- documento


def _portada(doc, cfg, anio):
    _parrafo(doc, INSTITUCION, negrita=True, tam=13, centrado=True)
    doc.add_paragraph()
    _parrafo(doc, cfg["titulo"], negrita=True, tam=16, centrado=True)
    _parrafo(doc, cfg["subtitulo"], negrita=True, tam=14, centrado=True)
    _parrafo(doc, str(anio), negrita=True, tam=14, centrado=True)
    doc.add_paragraph()
    _parrafo(
        doc,
        "Unidad de Planificación y Gestión de la Calidad",
        tam=11,
        centrado=True,
    )
    _parrafo(
        doc,
        "Generado el " + datetime.date.today().strftime("%d/%m/%Y"),
        tam=10,
        centrado=True,
    )
    doc.add_page_break()


def _preliminares(doc, cfg, anio):
    _titulo(doc, "1. Antecedentes")
    _parrafo(
        doc,
        f"En el año {anio} {cfg['sujeto']} del {INSTITUCION.title()} elaboraron su Plan "
        "Operativo Anual (POA), documento en el que constan las actividades principales que "
        f"la institución se propuso ejecutar durante el año {anio}.",
    )
    _titulo(doc, "2. Justificación")
    _parrafo(
        doc,
        "El Estatuto que rige las actividades académicas y administrativas de la institución "
        "solicita que las comisiones, direcciones de carrera, jefaturas de área y unidades "
        "emitan un informe del avance del trabajo planificado. El presente informe cuantifica "
        "el avance registrado hasta la fecha de su presentación.",
    )
    _titulo(doc, "3. Objetivos")
    _titulo(doc, "3.1. General", nivel=2)
    _parrafo(
        doc,
        "Socializar a las autoridades de la institución el avance general de las actividades "
        f"planificadas por {cfg['sujeto']}.",
    )
    _titulo(doc, "3.2. Específicos", nivel=2)
    for item in (
        f"Analizar los logros obtenidos por cada una de {cfg['sujeto']} durante el año {anio}.",
        "Analizar las diferencias fundamentales entre las actividades planificadas y "
        "ejecutadas en el periodo evaluado.",
        "Identificar las observaciones y evidencias reportadas por cada responsable.",
    ):
        doc.add_paragraph(item, style="List Bullet")


def _avance_general(doc, unidades, etiqueta, numero):
    _titulo(doc, f"{numero}. Avance general de {etiqueta}")
    promedio = sum(u.get("avance") or 0 for u in unidades) / len(unidades) if unidades else 0
    _parrafo(
        doc,
        f"El avance general de {etiqueta} es del {pct(promedio)} con respecto al plan "
        "operativo anual evaluado.",
    )
    filas = [
        [
            str(i),
            _nombre_unidad(u),
            u.get("responsable") or "Por definir",
            str(len(_actividades_validas(u))),
            pct(u.get("avance")),
        ]
        for i, u in enumerate(unidades, start=1)
    ]
    filas.append(["", "AVANCE GENERAL", "", "", pct(promedio)])
    _tabla(
        doc,
        ["Nro.", "Unidad / Carrera", "Responsable", "Actividades", "% Avance"],
        filas,
        anchos=[0.5, 2.8, 2.0, 0.9, 0.9],
    )
    _parrafo(doc, f"Tabla. Avance general de {etiqueta}", tam=9, centrado=True)

    grafico = _grafico_barras(
        f"Porcentaje de avance de {etiqueta}",
        [_nombre_unidad(u) for u in unidades],
        [u.get("avance") or 0 for u in unidades],
    )
    if grafico is not None:
        doc.add_picture(grafico, width=Inches(6.3))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _parrafo(doc, f"Figura. Porcentaje de avance de {etiqueta}", tam=9, centrado=True)
    return promedio


def _detalle_unidad(doc, unidad, indice, numero):
    _titulo(doc, f"{numero}.{indice} {_nombre_unidad(unidad)}", nivel=2)
    _parrafo(doc, f"Responsable: {unidad.get('responsable') or 'Por definir'}", negrita=True)
    _parrafo(doc, f"Categoría: {unidad.get('categoria') or '-'}", negrita=True)

    actividades = _actividades_validas(unidad)
    objetivos = []
    for a in actividades:
        obj = (a.get("objetivo") or "").strip()
        if obj and obj not in objetivos:
            objetivos.append(obj)
    if objetivos:
        _parrafo(doc, "Objetivos operativos:", negrita=True)
        for obj in objetivos:
            doc.add_paragraph(obj, style="List Bullet")

    _parrafo(
        doc,
        "Resumen del porcentaje de ejecución de las actividades programadas",
        negrita=True,
    )
    filas = [
        [str(i), a.get("tarea") or "-", pct(a.get("avance"))]
        for i, a in enumerate(actividades, start=1)
    ]
    filas.append(["", "TOTAL", pct(unidad.get("avance"))])
    _tabla(doc, ["Nro.", "Actividades planificadas", "% Ejecutado"], filas,
           anchos=[0.5, 5.3, 1.1])

    est = _estados(actividades)
    _parrafo(
        doc,
        f"De un total de {est['total']} actividades planificadas, {est['completadas']} se "
        f"encuentran ejecutadas al 100%, {est['proceso']} están en proceso y "
        f"{est['sin_iniciar']} no registran avance.",
    )

    _parrafo(doc, "Evidencias y observaciones", negrita=True)
    filas_obs = [
        [
            str(i),
            a.get("entregable") or "-",
            a.get("observacion") or "Sin observaciones",
        ]
        for i, a in enumerate(actividades, start=1)
    ]
    _tabla(doc, ["Nro.", "Evidencia / Entregable", "Observaciones"], filas_obs,
           anchos=[0.5, 3.2, 3.2])
    doc.add_page_break()


def _proyectos(doc, proyectos, numero):
    _titulo(doc, f"{numero}. Proyectos institucionales")
    if not proyectos:
        _parrafo(doc, "No se registran proyectos institucionales en el periodo evaluado.")
        return
    promedio = sum(p.get("avance") or 0 for p in proyectos) / len(proyectos)
    _parrafo(
        doc,
        f"La institución ejecuta {len(proyectos)} proyectos con un avance promedio del "
        f"{pct(promedio)}.",
    )
    filas = [
        [
            str(i),
            p.get("nombre") or "-",
            p.get("responsable") or "-",
            pct(p.get("avance")),
        ]
        for i, p in enumerate(proyectos, start=1)
    ]
    _tabla(doc, ["Nro.", "Proyecto", "Responsables", "% Avance"], filas,
           anchos=[0.5, 3.2, 2.5, 0.9])
    for i, p in enumerate(proyectos, start=1):
        _parrafo(
            doc,
            f"Proyecto {i}: {p.get('nombre') or '-'}. {p.get('descripcion') or ''}",
        )


def _cierre(doc, unidades, promedio, etiqueta, numero):
    actividades = [a for u in unidades for a in _actividades_validas(u)]
    est = _estados(actividades)
    _titulo(doc, f"{numero}. Conclusiones")
    for item in (
        f"El avance promedio de {etiqueta} alcanza el {pct(promedio)} del plan operativo anual "
        "evaluado.",
        f"Se registran {est['total']} actividades planificadas, de las cuales "
        f"{est['completadas']} están concluidas, {est['proceso']} en ejecución y "
        f"{est['sin_iniciar']} sin iniciar.",
        "Las unidades con menor avance requieren acompañamiento para cumplir la planificación "
        "en el periodo restante.",
    ):
        doc.add_paragraph(item, style="List Bullet")

    _titulo(doc, f"{numero + 1}. Recomendaciones")
    bajos = sorted(unidades, key=lambda u: u.get("avance") or 0)[:3]
    for item in (
        "Reforzar el seguimiento mensual de las actividades sin avance registrado.",
        "Solicitar a cada responsable la carga oportuna de las evidencias de cumplimiento.",
        "Priorizar el acompañamiento a: "
        + ", ".join(f"{_nombre_unidad(u)} ({pct(u.get('avance'))})" for u in bajos)
        + ".",
    ):
        doc.add_paragraph(item, style="List Bullet")


def generar_informe(poa: dict, tipo: str) -> tuple[BytesIO, str]:
    cfg = REPORTES[tipo]
    anio = poa.get("anio", datetime.date.today().year)
    unidades = [u for u in poa.get("unidades", []) if u.get("categoria") in cfg["categorias"]]

    doc = Document()
    seccion = doc.sections[0]
    seccion.orientation = WD_ORIENT.PORTRAIT
    seccion.left_margin = seccion.right_margin = Inches(0.9)
    estilo = doc.styles["Normal"]
    estilo.font.name = "Arial"
    estilo.font.size = Pt(11)

    pie = seccion.footer.paragraphs[0]
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pie.text = f"{cfg['subtitulo']} POA {anio} · {INSTITUCION}"
    pie.runs[0].font.size = Pt(7)

    _portada(doc, cfg, anio)
    _preliminares(doc, cfg, anio)

    numero = 4
    if tipo == "institucional":
        promedios = []
        for categoria in cfg["categorias"]:
            grupo = [u for u in unidades if u.get("categoria") == categoria]
            if not grupo:
                continue
            etiqueta = {
                "Unidad Administrativa": "las unidades administrativas",
                "Unidad Académica": "las unidades académicas",
                "Carrera": "las carreras",
            }[categoria]
            promedios.append(_avance_general(doc, grupo, etiqueta, numero))
            doc.add_page_break()
            numero += 1
        promedio = sum(promedios) / len(promedios) if promedios else 0
        _proyectos(doc, poa.get("proyectos", []), numero)
        numero += 1
        doc.add_page_break()
        _titulo(doc, f"{numero}. Avance individual por unidad")
        for i, u in enumerate(unidades, start=1):
            _detalle_unidad(doc, u, i, numero)
        numero += 1
        _cierre(doc, unidades, promedio, "la institución", numero)
    else:
        etiqueta = cfg["sujeto"]
        promedio = _avance_general(doc, unidades, etiqueta, numero)
        doc.add_page_break()
        _titulo(doc, f"{numero + 1}. Avance individual")
        for i, u in enumerate(unidades, start=1):
            _detalle_unidad(doc, u, i, numero + 1)
        _cierre(doc, unidades, promedio, etiqueta, numero + 2)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer, f"{cfg['archivo']}_{anio}.docx"
